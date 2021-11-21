import logging
import os
import re
from pathlib import Path
from typing import List, Optional

import sys
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))
from config import DOC_EXTENSIONS, DOCS_DIR
from core.models import DocumentChunk

log = logging.getLogger(__name__)

try:
    from pdfminer.high_level import extract_text as pdf_extract
    _HAS_PDF = True
except ImportError:
    _HAS_PDF = False
    log.warning("pdfminer.six not installed – PDF ingestion disabled.")

try:
    from docx import Document as DocxDocument
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False
    log.warning("python-docx not installed – DOCX ingestion disabled.")

try:
    from bs4 import BeautifulSoup
    _HAS_BS4 = True
except ImportError:
    _HAS_BS4 = False

try:
    import markdown as _md_lib
    _HAS_MD = True
except ImportError:
    _HAS_MD = False


# ─────────────────────────────────────────────────────────────────────────────
# Cleaning utilities
# ─────────────────────────────────────────────────────────────────────────────

_WS_RE    = re.compile(r"\s+")
_PUNCT_RE = re.compile(r"[^A-Za-z0-9\s\-_./]")


def _clean(text: str) -> str:
    text = _PUNCT_RE.sub(" ", text)
    return _WS_RE.sub(" ", text).strip()


def _chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
    """
    Split long text into overlapping word-chunks for embedding.
    chunk_size = 500 words is a reasonable balance for 2021-era Word2Vec models.
    """
    words  = text.split()
    chunks = []
    start  = 0
    while start < len(words):
        end = min(start + chunk_size, len(words))
        chunks.append(" ".join(words[start:end]))
        start += chunk_size - overlap
    return [c for c in chunks if len(c.split()) > 10]


# ─────────────────────────────────────────────────────────────────────────────
# Format-specific extractors
# ─────────────────────────────────────────────────────────────────────────────

def _extract_pdf(path: str) -> str:
    if not _HAS_PDF:
        return ""
    try:
        return pdf_extract(path) or ""
    except Exception as e:
        log.warning("PDF extraction failed (%s): %s", path, e)
        return ""


def _extract_docx(path: str) -> str:
    if not _HAS_DOCX:
        return ""
    try:
        doc   = DocxDocument(path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also grab table cell text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return " ".join(parts)
    except Exception as e:
        log.warning("DOCX extraction failed (%s): %s", path, e)
        return ""


def _extract_markdown(path: str) -> str:
    raw = Path(path).read_text(encoding="utf-8", errors="ignore")
    if _HAS_MD and _HAS_BS4:
        html = _md_lib.markdown(raw)
        return BeautifulSoup(html, "html.parser").get_text(separator=" ")
    # Fallback: strip markdown syntax manually
    raw = re.sub(r"#{1,6}\s+", "", raw)         # headings
    raw = re.sub(r"\*{1,2}([^*]+)\*{1,2}", r"\1", raw)  # bold/italic
    raw = re.sub(r"`{1,3}[^`]*`{1,3}", " ", raw)        # code spans
    raw = re.sub(r"!\[.*?\]\(.*?\)", " ", raw)           # images
    raw = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", raw)  # links
    return raw


def _extract_html(path: str) -> str:
    raw = Path(path).read_text(encoding="utf-8", errors="ignore")
    if _HAS_BS4:
        soup = BeautifulSoup(raw, "html.parser")
        for tag in soup(["script", "style", "meta", "head"]):
            tag.decompose()
        return soup.get_text(separator=" ")
    return re.sub(r"<[^>]+>", " ", raw)


def _extract_txt(path: str) -> str:
    return Path(path).read_text(encoding="utf-8", errors="ignore")


_EXTRACTORS = {
    ".pdf":  _extract_pdf,
    ".docx": _extract_docx,
    ".md":   _extract_markdown,
    ".html": _extract_html,
    ".htm":  _extract_html,
    ".txt":  _extract_txt,
}


# ─────────────────────────────────────────────────────────────────────────────
# Doc-type heuristic
# ─────────────────────────────────────────────────────────────────────────────

_DOC_TYPE_SIGNALS = {
    "requirements":  re.compile(r"requirement|shall|must|user.stor|acceptance", re.I),
    "architecture":  re.compile(r"architect|component|service|module|layer|design", re.I),
    "api-spec":      re.compile(r"endpoint|REST|GraphQL|swagger|openapi|grpc", re.I),
    "data-model":    re.compile(r"entity|table|schema|ERD|data.model|relation", re.I),
    "runbook":       re.compile(r"runbook|playbook|incident|deploy|rollback", re.I),
}


def _classify_doc_type(text: str, filename: str) -> str:
    sample = (filename + " " + text[:500]).lower()
    for dtype, pat in _DOC_TYPE_SIGNALS.items():
        if pat.search(sample):
            return dtype
    return "general"


# ─────────────────────────────────────────────────────────────────────────────
# Main ingester class
# ─────────────────────────────────────────────────────────────────────────────

class DocumentIngester:
    """
    Scans a folder (default: docs_input/) for supported document types
    and produces cleaned DocumentChunk objects.

    Also accepts explicit file paths via ingest_files().
    """

    def __init__(self, docs_folder: Optional[str] = None):
        self._docs_folder = docs_folder or str(DOCS_DIR)

    # ── Public ────────────────────────────────────────────────────────────────

    def ingest_folder(self) -> List[DocumentChunk]:
        """Scan the configured docs folder and return all chunks."""
        paths = []
        for fname in os.listdir(self._docs_folder):
            fpath = os.path.join(self._docs_folder, fname)
            if Path(fpath).is_file() and Path(fname).suffix.lower() in DOC_EXTENSIONS:
                paths.append(fpath)
        log.info("Found %d documents in %s", len(paths), self._docs_folder)
        return self.ingest_files(paths)

    def ingest_files(self, file_paths: List[str]) -> List[DocumentChunk]:
        """Ingest an explicit list of file paths."""
        chunks: List[DocumentChunk] = []
        for fp in file_paths:
            try:
                new_chunks = self._process_file(fp)
                chunks.extend(new_chunks)
                log.info("Ingested %d chunks from %s", len(new_chunks), fp)
            except Exception as e:
                log.error("Failed to ingest %s: %s", fp, e)
        return chunks

    def get_corpus_texts(self, chunks: List[DocumentChunk]) -> List[str]:
        """Return plain texts suitable for Word2Vec / TF-IDF augmentation."""
        return [c.text for c in chunks if c.text.strip()]

    # ── Internals ─────────────────────────────────────────────────────────────

    def _process_file(self, path: str) -> List[DocumentChunk]:
        ext       = Path(path).suffix.lower()
        extractor = _EXTRACTORS.get(ext)
        if not extractor:
            log.warning("No extractor for extension %s (%s)", ext, path)
            return []

        raw_text  = extractor(path)
        if not raw_text or len(raw_text.strip()) < 50:
            log.debug("Empty or too-short extraction from %s", path)
            return []

        cleaned   = _clean(raw_text)
        doc_type  = _classify_doc_type(cleaned, Path(path).name)
        text_chunks = _chunk_text(cleaned)

        return [
            DocumentChunk(
                source_file=path,
                text=chunk,
                doc_type=doc_type,
            )
            for chunk in text_chunks
        ]
